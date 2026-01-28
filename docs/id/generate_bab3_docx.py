#!/usr/bin/env python3
"""
Generate BAB 3 METODOLOGI PENELITIAN DOCX
Mengikuti format BINUS Online sesuai Petunjuk Penulisan Skripsi

Format:
- Kertas: A4 (80gr)
- Margin: Kiri 4cm, Atas/Bawah/Kanan 2.5cm
- Font: Times New Roman 12pt
- Spasi: 1.5
- Judul Bab: 12pt, Bold, KAPITAL, Center
- Judul Subbab: 12pt, Bold, Title Case, Left
- Gambar: Judul di bawah, center
- Tabel: Judul di atas, center
- Istilah asing: Italic
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Paths
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "BAB_3_METODOLOGI_PENELITIAN.docx")
PICS_DIR = os.path.join(SCRIPT_DIR, "diagrams", "pics")

# Image mapping
IMAGE_FILES = {
    "3.1": "BAB_3_Gambar_3.1_DSR_Framework.drawio.drawio.drawio.png",
    "3.2": "BAB_3_Gambar_3.2_Matriks_Variabel_Keputusan.drawio.png",
    "3.3": "BAB_3_Gambar_3.3_Hierarki_Constraint.drawio.png",
    "3.4": "BAB_3_Gambar_3.4_Arsitektur_Solver.drawio.png",
    "3.5": "BAB_3_Gambar_3.5_Arsitektur_Sistem.drawio.png",
    "3.6": "BAB_3_Gambar_3.6_Matriks_Pengumpulan_Data.drawio.png",
    "3.7": "BAB_3_Gambar_3.7_Flowchart_Optimisasi.drawio.png",
    "3.8": "BAB_3_Gambar_3.8_Piramida_Strategi_Pengujian.drawio.png",
}


def set_cell_margins(cell, top=0, bottom=0, left=108, right=108):
    """Set cell margins in twips (1/20 of a point)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, margin_value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(margin_value))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_table_borders(table):
    """Set table borders to single black lines"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)

    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def create_document():
    """Create document with BINUS formatting"""
    doc = Document()

    # Set page margins: Left 4cm, Top/Bottom/Right 2.5cm
    for section in doc.sections:
        section.left_margin = Cm(4)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.page_width = Cm(21)  # A4
        section.page_height = Cm(29.7)  # A4

    # Create styles
    create_styles(doc)

    return doc


def create_styles(doc):
    """Create custom styles for BINUS format"""
    styles = doc.styles

    # Normal style - Times New Roman 12pt, 1.5 spacing
    normal = styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    # Chapter Title - Bold, CAPITALS, Center
    if 'ChapterTitle' not in [s.name for s in styles]:
        chapter_style = styles.add_style('ChapterTitle', WD_STYLE_TYPE.PARAGRAPH)
        chapter_style.font.name = 'Times New Roman'
        chapter_style.font.size = Pt(12)
        chapter_style.font.bold = True
        chapter_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chapter_style.paragraph_format.space_before = Pt(0)
        chapter_style.paragraph_format.space_after = Pt(24)
        chapter_style.paragraph_format.line_spacing = 1.5

    # SubSection Title (3.1, 3.2, etc.) - Bold, Left
    if 'SubSectionTitle' not in [s.name for s in styles]:
        sub_style = styles.add_style('SubSectionTitle', WD_STYLE_TYPE.PARAGRAPH)
        sub_style.font.name = 'Times New Roman'
        sub_style.font.size = Pt(12)
        sub_style.font.bold = True
        sub_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sub_style.paragraph_format.space_before = Pt(18)
        sub_style.paragraph_format.space_after = Pt(12)
        sub_style.paragraph_format.line_spacing = 1.5

    # SubSubSection Title (3.1.1, 3.1.2, etc.) - Bold, Left
    if 'SubSubSectionTitle' not in [s.name for s in styles]:
        subsub_style = styles.add_style('SubSubSectionTitle', WD_STYLE_TYPE.PARAGRAPH)
        subsub_style.font.name = 'Times New Roman'
        subsub_style.font.size = Pt(12)
        subsub_style.font.bold = True
        subsub_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        subsub_style.paragraph_format.space_before = Pt(12)
        subsub_style.paragraph_format.space_after = Pt(6)
        subsub_style.paragraph_format.line_spacing = 1.5

    # Figure Caption - Center, below figure
    if 'FigureCaption' not in [s.name for s in styles]:
        fig_style = styles.add_style('FigureCaption', WD_STYLE_TYPE.PARAGRAPH)
        fig_style.font.name = 'Times New Roman'
        fig_style.font.size = Pt(12)
        fig_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fig_style.paragraph_format.space_before = Pt(6)
        fig_style.paragraph_format.space_after = Pt(12)
        fig_style.paragraph_format.line_spacing = 1.5

    # Table Caption - Center, above table
    if 'TableCaption' not in [s.name for s in styles]:
        table_style = styles.add_style('TableCaption', WD_STYLE_TYPE.PARAGRAPH)
        table_style.font.name = 'Times New Roman'
        table_style.font.size = Pt(12)
        table_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table_style.paragraph_format.space_before = Pt(12)
        table_style.paragraph_format.space_after = Pt(6)
        table_style.paragraph_format.line_spacing = 1.5


def add_paragraph(doc, text, style='Normal', italic_terms=None):
    """Add paragraph with optional italic terms"""
    p = doc.add_paragraph(style=style)

    if italic_terms:
        # Split text and apply italic to specific terms
        remaining = text
        for term in italic_terms:
            if term in remaining:
                parts = remaining.split(term, 1)
                if parts[0]:
                    run = p.add_run(parts[0])
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                run = p.add_run(term)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.italic = True
                remaining = parts[1] if len(parts) > 1 else ""
        if remaining:
            run = p.add_run(remaining)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    else:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    return p


def add_bold_paragraph(doc, text, style='Normal'):
    """Add bold paragraph"""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    return p


def add_image(doc, image_key, caption, width_inches=5.5):
    """Add image with caption below"""
    img_filename = IMAGE_FILES.get(image_key)
    if img_filename:
        img_path = os.path.join(PICS_DIR, img_filename)
        if os.path.exists(img_path):
            # Add image centered
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Inches(width_inches))

            # Add caption below
            cap = doc.add_paragraph(caption, style='FigureCaption')
            return True

    # Fallback if image not found
    p = doc.add_paragraph(f"[Gambar tidak ditemukan: {img_filename}]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption, style='FigureCaption')
    return False


def add_table(doc, headers, rows, caption=None):
    """Add table with caption above"""
    # Caption above table
    if caption:
        doc.add_paragraph(caption, style='TableCaption')

    # Create table
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.bold = True

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data)
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)

    # Add space after table
    doc.add_paragraph()

    return table


def add_code_block(doc, code, language="python"):
    """Add code block with monospace font"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    return p


def add_bullet_list(doc, items):
    """Add bullet list"""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)


def add_numbered_list(doc, items):
    """Add numbered list"""
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {item}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)


def build_document():
    """Build the complete BAB 3 document"""
    doc = create_document()

    # ==========================================
    # CHAPTER TITLE
    # ==========================================
    doc.add_paragraph("BAB 3", style='ChapterTitle')
    p = doc.add_paragraph("METODOLOGI PENELITIAN", style='ChapterTitle')

    # ==========================================
    # 3.1 Kerangka Penelitian
    # ==========================================
    doc.add_paragraph("3.1 Kerangka Penelitian", style='SubSectionTitle')

    add_paragraph(doc,
        "Penelitian ini mengadopsi Design Science Research (DSR) methodology yang dikembangkan oleh "
        "Hevner et al. (2004) dan Peffers et al. (2007). DSR dipilih karena fokus penelitian adalah "
        "pengembangan artefak teknologi (sistem penjadwalan) yang memberikan solusi praktis untuk "
        "permasalahan nyata di industri perhotelan.",
        italic_terms=["Design Science Research", "methodology"]
    )

    doc.add_paragraph("3.1.1 Design Science Research Framework", style='SubSubSectionTitle')

    add_image(doc, "3.1", "Gambar 3.1: Design Science Research Framework (Adaptasi dari Hevner et al., 2004)")

    doc.add_paragraph("3.1.2 Tahapan Penelitian DSR", style='SubSubSectionTitle')

    add_table(doc,
        ["Tahap", "Aktivitas", "Output"],
        [
            ["1. Problem Identification", "Analisis permasalahan penjadwalan shift di divisi dapur hotel", "Problem statement & requirements"],
            ["2. Define Objectives", "Mendefinisikan kriteria solusi optimal", "Success metrics & KPIs"],
            ["3. Design & Development", "Pengembangan model CP-SAT dengan soft constraints", "Working prototype"],
            ["4. Demonstration", "Implementasi pada studi kasus nyata", "Deployed system"],
            ["5. Evaluation", "Pengujian dan analisis performa", "Evaluation report"],
            ["6. Communication", "Dokumentasi dan publikasi hasil", "Thesis & publications"],
        ],
        caption="Tabel 3.1: Tahapan Penelitian DSR"
    )

    # ==========================================
    # 3.2 Pemodelan Constraint Satisfaction Problem (CSP)
    # ==========================================
    doc.add_paragraph("3.2 Pemodelan Constraint Satisfaction Problem (CSP)", style='SubSectionTitle')

    doc.add_paragraph("3.2.1 Definisi Formal CSP", style='SubSubSectionTitle')

    add_paragraph(doc,
        "Berdasarkan Russell & Norvig (2020), Constraint Satisfaction Problem (CSP) didefinisikan "
        "sebagai tuple (X, D, C) dimana:",
        italic_terms=["Constraint Satisfaction Problem"]
    )

    add_bullet_list(doc, [
        "X = {x1, x2, ..., xn} adalah himpunan variabel",
        "D = {D1, D2, ..., Dn} adalah domain untuk setiap variabel",
        "C = {C1, C2, ..., Cm} adalah himpunan constraint"
    ])

    doc.add_paragraph("3.2.2 Formulasi CSP untuk Penjadwalan Shift", style='SubSubSectionTitle')

    add_paragraph(doc,
        "Dalam konteks penjadwalan shift divisi dapur hotel, CSP dirumuskan sebagai berikut:"
    )

    add_bold_paragraph(doc, "Variabel Keputusan:")
    add_paragraph(doc, "X = {shifts[s,d,t] | s E S, d E D, t E T}")
    add_paragraph(doc, "Dimana:")
    add_bullet_list(doc, [
        "S = Himpunan staf (s1, s2, ..., sn)",
        "D = Himpunan tanggal dalam periode (d1, d2, ..., dm)",
        "T = Himpunan tipe shift {WORK, OFF, EARLY, LATE}"
    ])

    add_bold_paragraph(doc, "Domain:")
    add_paragraph(doc, "D(shifts[s,d,t]) = {0, 1} (Boolean)")
    add_paragraph(doc, "Nilai 1 = Staf s mendapat shift tipe t pada tanggal d")
    add_paragraph(doc, "Nilai 0 = Staf s TIDAK mendapat shift tipe t pada tanggal d")

    add_bold_paragraph(doc, "Total Variabel:")
    add_paragraph(doc, "|X| = |S| x |D| x |T|")
    add_paragraph(doc, "Contoh: 15 staf x 60 hari x 4 tipe = 3,600 variabel boolean")

    doc.add_paragraph("3.2.3 Visualisasi Model Variabel", style='SubSubSectionTitle')

    add_image(doc, "3.2", "Gambar 3.2: Struktur Matriks Variabel Keputusan")

    # ==========================================
    # 3.3 Klasifikasi Constraint: Hard vs Soft
    # ==========================================
    doc.add_paragraph("3.3 Klasifikasi Constraint: Hard vs Soft", style='SubSectionTitle')

    doc.add_paragraph("3.3.1 Pendekatan Soft Constraint Optimization", style='SubSubSectionTitle')

    add_paragraph(doc,
        "Berbeda dengan sistem penjadwalan tradisional yang menggunakan hard constraints "
        "(harus dipenuhi 100%), penelitian ini mengadopsi Soft Constraint Optimization approach "
        "yang terinspirasi dari karya Verfaillie & Jussien (2005) dan penelitian Rossi et al. (2006).",
        italic_terms=["hard constraints", "Soft Constraint Optimization"]
    )

    add_bold_paragraph(doc, "Keunggulan Soft Constraint:")
    add_numbered_list(doc, [
        "Always-Feasible Solutions: Selalu menghasilkan solusi, tidak pernah INFEASIBLE",
        "Trade-off Optimization: Memungkinkan kompromi antar constraint yang berkonflik",
        "Real-world Applicability: Lebih sesuai dengan kondisi nyata operasional hotel"
    ])

    doc.add_paragraph("3.3.2 Tabel Klasifikasi Constraint", style='SubSubSectionTitle')

    add_table(doc,
        ["No", "Constraint", "Tipe", "Deskripsi", "Penalti"],
        [
            ["1", "One Shift Per Day", "HARD", "Setiap staf hanya 1 tipe shift per hari", "-"],
            ["2", "Pre-filled Cells", "HARD", "Sel yang sudah diisi manager tidak berubah", "-"],
            ["3", "Calendar Must-Off", "HARD", "Tanggal wajib libur (hari raya)", "-"],
            ["4", "Staff Group", "SOFT", "Maksimal 1 anggota off/early per grup per hari", "100"],
            ["5", "Daily Limit Min", "SOFT", "Minimal staf off per hari", "50"],
            ["6", "Daily Limit Max", "SOFT", "Maksimal staf off per hari", "50"],
            ["7", "Monthly Limit", "SOFT", "Min/max hari libur per staf per periode", "80"],
            ["8", "5-Day Rest", "SOFT", "Maksimal 5 hari kerja berturut-turut", "200"],
            ["9", "Staff Type Limit", "SOFT", "Batas per tipe staf", "60"],
            ["10", "Adjacent Conflict", "SOFT", "Hindari pola berurutan tertentu", "30"],
        ],
        caption="Tabel 3.2: Klasifikasi Constraint"
    )

    doc.add_paragraph("3.3.3 Diagram Hierarki Constraint", style='SubSubSectionTitle')

    add_image(doc, "3.3", "Gambar 3.3: Hierarki Constraint berdasarkan Prioritas")

    # ==========================================
    # 3.4 Formulasi Objective Function
    # ==========================================
    doc.add_paragraph("3.4 Formulasi Objective Function", style='SubSectionTitle')

    doc.add_paragraph("3.4.1 Fungsi Objektif", style='SubSubSectionTitle')

    add_paragraph(doc,
        "Berdasarkan pendekatan penalty-based optimization dari Hooker (2007), "
        "fungsi objektif dirumuskan sebagai minimalisasi total penalti:",
        italic_terms=["penalty-based optimization"]
    )

    add_paragraph(doc, "Minimize: Z = SUM(wi x vi) untuk semua i dalam V")
    add_paragraph(doc, "Dimana:")
    add_bullet_list(doc, [
        "V = Himpunan soft constraint violations",
        "wi = Bobot penalti untuk violation i",
        "vi = Variabel boolean (1 jika terjadi violation, 0 jika tidak)"
    ])

    doc.add_paragraph("3.4.2 Implementasi dalam OR-Tools CP-SAT", style='SubSubSectionTitle')

    add_paragraph(doc,
        "Implementasi fungsi objektif dalam Python menggunakan OR-Tools CP-SAT:"
    )

    code = '''def _add_objective(self):
    """
    Build the objective function: Minimize total weighted violations.
    """
    if not self.violation_vars:
        return

    objective_terms = []
    for violation_var, weight, description in self.violation_vars:
        objective_terms.append(violation_var * weight)

    self.model.Minimize(sum(objective_terms))'''

    add_code_block(doc, code)

    doc.add_paragraph("3.4.3 Konfigurasi Default Penalty Weights", style='SubSubSectionTitle')

    add_table(doc,
        ["Constraint", "Bobot Penalti", "Prioritas"],
        [
            ["staff_group", "100", "Tinggi - coverage grup penting"],
            ["daily_limit", "50", "Medium - keseimbangan harian"],
            ["daily_limit_max", "50", "Medium - batas maksimal harian"],
            ["monthly_limit", "80", "Tinggi - keadilan bulanan"],
            ["adjacent_conflict", "30", "Rendah - kenyamanan"],
            ["5_day_rest", "200", "Sangat tinggi - kepatuhan regulasi"],
            ["staff_type_limit", "60", "Medium-tinggi - coverage per tipe"],
            ["backup_coverage", "500", "Tertinggi - kontinuitas operasional"],
        ],
        caption="Tabel 3.3: Konfigurasi Default Penalty Weights"
    )

    # ==========================================
    # 3.5 Arsitektur Solver CP-SAT
    # ==========================================
    doc.add_paragraph("3.5 Arsitektur Solver CP-SAT", style='SubSectionTitle')

    doc.add_paragraph("3.5.1 Tentang Google OR-Tools CP-SAT", style='SubSubSectionTitle')

    add_paragraph(doc,
        "CP-SAT (Constraint Programming - Satisfiability) adalah constraint solver hybrid "
        "yang dikembangkan oleh Google. Menurut dokumentasi resmi Google OR-Tools (2024) "
        "dan Perron & Furnon (2023), CP-SAT menggabungkan:",
        italic_terms=["Constraint Programming", "Satisfiability", "hybrid"]
    )

    add_numbered_list(doc, [
        "Constraint Programming (CP) - Domain propagation dan inference",
        "SAT Solving - Boolean satisfiability techniques",
        "Linear Programming (LP) - Relaxation dan cutting planes",
        "Local Search - Large Neighborhood Search (LNS)"
    ])

    doc.add_paragraph("3.5.2 Diagram Arsitektur Solver", style='SubSubSectionTitle')

    add_image(doc, "3.4", "Gambar 3.4: Arsitektur Google OR-Tools CP-SAT Solver")

    doc.add_paragraph("3.5.3 Konfigurasi Solver", style='SubSubSectionTitle')

    add_table(doc,
        ["Parameter", "Nilai Default", "Deskripsi"],
        [
            ["max_time_in_seconds", "30", "Timeout maksimal pencarian"],
            ["num_search_workers", "4", "Jumlah worker paralel"],
            ["log_search_progress", "True", "Logging progress pencarian"],
        ],
        caption="Tabel 3.4: Konfigurasi Solver CP-SAT"
    )

    # ==========================================
    # 3.6 Arsitektur Sistem Keseluruhan
    # ==========================================
    doc.add_paragraph("3.6 Arsitektur Sistem Keseluruhan", style='SubSectionTitle')

    doc.add_paragraph("3.6.1 Diagram Arsitektur Sistem", style='SubSubSectionTitle')

    add_paragraph(doc,
        "Sistem penjadwalan shift terdiri dari empat lapisan utama yang terintegrasi: "
        "User Interface Layer (React + Tailwind CSS), Application Layer (React Hooks + State Management), "
        "Optimization Service Layer (Python Flask + OR-Tools), dan Data Layer (Supabase PostgreSQL)."
    )

    add_image(doc, "3.5", "Gambar 3.5: Arsitektur Sistem Penjadwalan Shift")

    # ==========================================
    # 3.7 Pengumpulan Data
    # ==========================================
    doc.add_paragraph("3.7 Pengumpulan Data", style='SubSectionTitle')

    doc.add_paragraph("3.7.1 Sumber Data", style='SubSubSectionTitle')

    add_paragraph(doc,
        "Data penelitian diperoleh dari studi kasus nyata divisi dapur hotel dengan karakteristik:"
    )

    add_table(doc,
        ["Aspek", "Detail"],
        [
            ["Lokasi", "Hotel bintang 4-5"],
            ["Divisi", "Kitchen/Dapur"],
            ["Jumlah Staf", "15-20 orang"],
            ["Periode Data", "Januari - Desember 2024"],
            ["Tipe Staf", "Tetap, Kontrak, Part-time"],
        ],
        caption="Tabel 3.5: Karakteristik Data Penelitian"
    )

    doc.add_paragraph("3.7.2 Jenis Data yang Dikumpulkan", style='SubSubSectionTitle')

    add_image(doc, "3.6", "Gambar 3.6: Matriks Pengumpulan Data")

    # ==========================================
    # 3.8 Metode Evaluasi
    # ==========================================
    doc.add_paragraph("3.8 Metode Evaluasi", style='SubSectionTitle')

    doc.add_paragraph("3.8.1 Metrik Evaluasi Kuantitatif", style='SubSubSectionTitle')

    add_table(doc,
        ["No", "Metrik", "Formula", "Target"],
        [
            ["1", "Constraint Satisfaction Rate", "(Total - Violations) / Total x 100%", ">= 95%"],
            ["2", "Hard Constraint Satisfaction", "Hard satisfied / Total hard x 100%", "100%"],
            ["3", "Soft Constraint Satisfaction", "Soft satisfied / Total soft x 100%", ">= 85%"],
            ["4", "Computation Time", "Waktu dari input hingga output", "<= 30s"],
            ["5", "Solution Quality Score", "1 - (Total penalty / Max penalty)", ">= 0.9"],
            ["6", "Fairness Index", "Std deviation hari libur antar staf", "<= 2 hari"],
        ],
        caption="Tabel 3.6: Metrik Evaluasi Kuantitatif"
    )

    doc.add_paragraph("3.8.2 Metrik Evaluasi Kualitatif", style='SubSubSectionTitle')

    add_table(doc,
        ["No", "Aspek", "Metode Pengukuran"],
        [
            ["1", "User Satisfaction", "Kuesioner skala Likert 1-5"],
            ["2", "Ease of Use", "System Usability Scale (SUS)"],
            ["3", "Perceived Fairness", "Interview dengan staf"],
            ["4", "Manager Acceptance", "Focus Group Discussion"],
        ],
        caption="Tabel 3.7: Metrik Evaluasi Kualitatif"
    )

    doc.add_paragraph("3.8.3 Perbandingan dengan Metode Baseline", style='SubSubSectionTitle')

    add_table(doc,
        ["Aspek", "Manual", "Rule-based", "CP-SAT (Proposed)"],
        [
            ["Waktu pembuatan", "4-8 jam", "5-10 menit", "< 30 detik"],
            ["Optimality", "Tidak terjamin", "Heuristic", "Mathematically optimal"],
            ["Constraint handling", "Trial-and-error", "Sequential", "Simultaneous"],
            ["Scalability", "Tidak scalable", "Terbatas", "Highly scalable"],
            ["Consistency", "Bervariasi", "Konsisten", "Optimal & consistent"],
        ],
        caption="Tabel 3.8: Perbandingan Metode Penjadwalan"
    )

    # ==========================================
    # 3.9 Implementasi Constraint dalam Kode
    # ==========================================
    doc.add_paragraph("3.9 Implementasi Constraint dalam Kode", style='SubSectionTitle')

    doc.add_paragraph("3.9.1 Basic Constraint: One Shift Per Day", style='SubSubSectionTitle')

    add_paragraph(doc,
        "Implementasi constraint dasar yang memastikan setiap staf hanya memiliki satu tipe shift per hari:"
    )

    code = '''def _add_basic_constraints(self):
    """Each staff has exactly one shift type per day."""
    for staff in self.staff_members:
        for date in self.date_range:
            self.model.AddExactlyOne([
                self.shifts[(staff['id'], date, shift)]
                for shift in range(4)  # WORK, OFF, EARLY, LATE
            ])'''

    add_code_block(doc, code)

    doc.add_paragraph("3.9.2 Soft Constraint: Staff Group", style='SubSubSectionTitle')

    add_paragraph(doc,
        "Implementasi soft constraint untuk grup staf dengan konsep off-equivalent:",
        italic_terms=["soft constraint", "off-equivalent"]
    )

    code = '''def _add_staff_group_constraints(self):
    """Maximum 1 member from each group can be off OR early per day."""
    for group in self.staff_groups:
        for date in self.date_range:
            violation = self.model.NewBoolVar(f'group_violation_{group["id"]}_{date}')

            off_equivalent = sum(
                2 * self.shifts[(s['id'], date, self.SHIFT_OFF)] +
                self.shifts[(s['id'], date, self.SHIFT_EARLY)]
                for s in group['members']
            )

            self.model.Add(off_equivalent <= 2).OnlyEnforceIf(violation.Not())

            self.violation_vars.append(
                (violation, self.PENALTY_WEIGHTS['staff_group'], f'group_{group["id"]}_{date}')
            )'''

    add_code_block(doc, code)

    doc.add_paragraph("3.9.3 Soft Constraint: 5-Day Rest", style='SubSubSectionTitle')

    add_paragraph(doc,
        "Implementasi constraint untuk memastikan maksimal 5 hari kerja berturut-turut:"
    )

    code = '''def _add_5_day_rest_constraint(self):
    """Maximum 5 consecutive work days."""
    for staff in self.staff_members:
        for i in range(len(self.date_range) - 5):
            violation = self.model.NewBoolVar(f'5day_{staff["id"]}_{i}')

            work_days = sum(
                self.shifts[(staff['id'], self.date_range[i+j], self.SHIFT_WORK)]
                for j in range(6)
            )

            self.model.Add(work_days <= 5).OnlyEnforceIf(violation.Not())

            self.violation_vars.append(
                (violation, self.PENALTY_WEIGHTS['5_day_rest'], f'5day_{staff["id"]}_{i}')
            )'''

    add_code_block(doc, code)

    # ==========================================
    # 3.10 Alur Proses Optimisasi
    # ==========================================
    doc.add_paragraph("3.10 Alur Proses Optimisasi", style='SubSectionTitle')

    doc.add_paragraph("3.10.1 Flowchart Proses Optimisasi", style='SubSubSectionTitle')

    add_paragraph(doc,
        "Proses optimisasi dimulai dari input data (staf, tanggal, constraint, pre-filled), "
        "dilanjutkan dengan inisialisasi model CP-SAT, pembuatan variabel keputusan, "
        "penambahan constraint (hard dan soft), pembangunan fungsi objektif, konfigurasi solver, "
        "eksekusi solver, dan ekstraksi solusi."
    )

    add_image(doc, "3.7", "Gambar 3.7: Flowchart Proses Optimisasi")

    # ==========================================
    # 3.11 Validasi dan Verifikasi
    # ==========================================
    doc.add_paragraph("3.11 Validasi dan Verifikasi", style='SubSectionTitle')

    doc.add_paragraph("3.11.1 Strategi Pengujian", style='SubSubSectionTitle')

    add_paragraph(doc,
        "Strategi pengujian mengikuti piramida pengujian yang terdiri dari Unit Tests "
        "(pengujian fungsi constraint individual), Integration Tests (pengujian API + Solver), "
        "dan End-to-End Tests (otomatisasi browser)."
    )

    add_image(doc, "3.8", "Gambar 3.8: Piramida Strategi Pengujian")

    doc.add_paragraph("3.11.2 Test Cases", style='SubSubSectionTitle')

    add_table(doc,
        ["ID", "Test Case", "Input", "Expected Output"],
        [
            ["TC01", "Basic constraint", "5 staff, 7 days", "Each staff has 1 shift/day"],
            ["TC02", "Pre-filled preservation", "Pre-filled x day 1", "x preserved in output"],
            ["TC03", "Staff group constraint", "2 members in group", "Max 1 off per day"],
            ["TC04", "5-day rest constraint", "7 consecutive days", "At least 1 rest in days 1-6"],
            ["TC05", "Monthly limit", "Min=8, Max=10", "Output has 8-10 off days"],
            ["TC06", "Large scale performance", "20 staff, 60 days", "Solution in < 30s"],
        ],
        caption="Tabel 3.9: Test Cases"
    )

    # ==========================================
    # 3.12 Jadwal Penelitian
    # ==========================================
    doc.add_paragraph("3.12 Jadwal Penelitian", style='SubSectionTitle')

    add_table(doc,
        ["Fase", "Aktivitas", "Durasi"],
        [
            ["Fase 1", "Studi literatur & analisis kebutuhan", "4 minggu"],
            ["Fase 2", "Desain model CSP & constraint mapping", "3 minggu"],
            ["Fase 3", "Implementasi OR-Tools optimizer", "4 minggu"],
            ["Fase 4", "Integrasi dengan UI & testing", "3 minggu"],
            ["Fase 5", "Evaluasi & analisis hasil", "2 minggu"],
            ["Fase 6", "Penulisan laporan & revisi", "4 minggu"],
        ],
        caption="Tabel 3.10: Jadwal Penelitian"
    )

    # ==========================================
    # 3.13 Referensi
    # ==========================================
    doc.add_paragraph("3.13 Referensi", style='SubSectionTitle')

    add_bold_paragraph(doc, "Referensi Akademis")

    references = [
        "Hevner, A. R., March, S. T., Park, J., & Ram, S. (2004). Design Science in Information Systems Research. MIS Quarterly, 28(1), 75-105.",
        "Peffers, K., Tuunanen, T., Rothenberger, M. A., & Chatterjee, S. (2007). A Design Science Research Methodology for Information Systems Research. Journal of Management Information Systems, 24(3), 45-77.",
        "Russell, S. J., & Norvig, P. (2020). Artificial Intelligence: A Modern Approach (4th ed.). Pearson.",
        "Rossi, F., Van Beek, P., & Walsh, T. (2006). Handbook of Constraint Programming. Elsevier.",
        "Verfaillie, G., & Jussien, N. (2005). Constraint Solving in Uncertain and Dynamic Environments: A Survey. Constraints, 10(3), 253-281.",
        "Hooker, J. N. (2007). Integrated Methods for Optimization. Springer.",
        "Krupke, D. (2024). CP-SAT Primer: Using and Understanding Google OR-Tools' CP-SAT Solver. TU Braunschweig.",
        "Van Den Bergh, J., Belien, J., De Bruecker, P., Demeulemeester, E., & De Boeck, L. (2013). Personnel Scheduling: A Literature Review. European Journal of Operational Research, 226(3), 367-385.",
        "Burke, E. K., De Causmaecker, P., Berghe, G. V., & Van Landeghem, H. (2004). The State of the Art of Nurse Rostering. Journal of Scheduling, 7(6), 441-499.",
        "Ernst, A. T., Jiang, H., Krishnamoorthy, M., & Sier, D. (2004). Staff Scheduling and Rostering: A Review of Applications, Methods and Models. European Journal of Operational Research, 153(1), 3-27.",
    ]

    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"[{i}] {ref}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

    add_bold_paragraph(doc, "Referensi Teknis")

    tech_refs = [
        "Google OR-Tools Documentation (2024). CP-SAT Solver Guide. https://developers.google.com/optimization/cp/cp_solver",
        "Google OR-Tools GitHub Repository (2024). Constraint Programming Examples. https://github.com/google/or-tools",
        "Perron, L., & Furnon, V. (2023). OR-Tools by Google. Operations Research Tools. https://developers.google.com/optimization",
        "Flask Documentation (2024). Flask Web Framework. https://flask.palletsprojects.com/",
        "React Documentation (2024). React: A JavaScript Library for Building User Interfaces. https://react.dev/",
    ]

    for i, ref in enumerate(tech_refs, 11):
        p = doc.add_paragraph()
        run = p.add_run(f"[{i}] {ref}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

    return doc


def main():
    """Main function to generate the document"""
    print("Generating BAB 3 METODOLOGI PENELITIAN DOCX...")
    print(f"Output: {OUTPUT_PATH}")
    print(f"Images directory: {PICS_DIR}")

    # Check if images exist
    print("\nChecking images:")
    for key, filename in IMAGE_FILES.items():
        path = os.path.join(PICS_DIR, filename)
        status = "OK" if os.path.exists(path) else "NOT FOUND"
        print(f"  Gambar {key}: {status}")

    # Build document
    doc = build_document()

    # Save
    doc.save(OUTPUT_PATH)
    print(f"\nDocument saved: {OUTPUT_PATH}")

    # Verify
    if os.path.exists(OUTPUT_PATH):
        size = os.path.getsize(OUTPUT_PATH)
        print(f"File size: {size:,} bytes")
        print("SUCCESS!")
    else:
        print("ERROR: File not created")


if __name__ == "__main__":
    main()
